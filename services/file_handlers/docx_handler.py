from docx import Document
from datetime import datetime
from typing import List, Dict, Any
from .base_handler import FileHandler, chunk_text
from io import BytesIO

class DocxHandler(FileHandler):
    def process(self, content: bytes, filename: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        document = Document(BytesIO(content))
        text_content = "\n".join([para.text for para in document.paragraphs])
        chunks = chunk_text(text_content, chunk_size=chunk_size, overlap=overlap)
        docx_metadata = self.get_metadata(content, filename)

        processed_chunks = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                **docx_metadata,
                'chunk_id': i,
                'total_chunks': len(chunks),
                'chunk_text_length': len(chunk),
                'chunk_word_count': len(chunk.split())
            }
            processed_chunks.append({'content': chunk, 'metadata': chunk_metadata})
        return processed_chunks

    def get_metadata(self, content: bytes, filename: str) -> Dict[str, Any]:
        document = Document(BytesIO(content))
        core_properties = document.core_properties
        return {
            'filename': filename,
            'file_type': 'docx',
            'file_size': len(content),
            'upload_timestamp': datetime.now().isoformat(),
            'author': core_properties.author,
            'created': core_properties.created,
            'modified': core_properties.modified,
            'title': core_properties.title,
        }
